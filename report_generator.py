from docx import Document
import pandas
import matplotlib.pyplot as plt
import report_tools
import tools
from bs4 import BeautifulSoup
from docx.shared import Inches
import traceback

#500012 for testing
security_codes=input("Enter the security codes seperated by commas: ").split(",")

for code in security_codes:
    report=Document()
    ticker,company_name=report_tools.security_code_to_ticker_and_name(code)
    report.add_heading(f"{company_name}({ticker})",level=0)
    with open("/home/yashas/Desktop/Quant/NEW_DATA/"+code+".html","r") as f:
        html_content=f.read()
    soup = BeautifulSoup(html_content, 'html.parser')
    #Some Numbers
    report.add_heading("Some Numbers",level=2)
    numbres=soup.findAll("li",class_="flex flex-space-between")
    for i in numbres:
        if "Market Cap" in i.text:
            market_cap=i.text
        if "Stock P/E" in i.text:
            stock_pe=i.text
    market_cap=market_cap.replace("\n","")
    market_cap=market_cap.replace("Market Cap","")
    market_cap=market_cap.replace("₹","")
    market_cap=market_cap.replace("Cr.","")
    market_cap=market_cap.replace(" ","")
    stock_pe=stock_pe.replace("\n","")
    stock_pe=stock_pe.replace("Stock P/E","")
    stock_pe=stock_pe.replace(" ","")
    report.add_paragraph(f"Market Cap: {market_cap}Cr, Stock P/E: {stock_pe}")
    #about
    about_element = soup.findAll('div', class_="sub show-more-box about")
    about=""
    for element in about_element:
        about=about+element.text
    about.replace("[1]","")
    report.add_heading("About",level=2)
    report.add_paragraph(about)
    #cons
    cons_element = soup.findAll('div', class_="cons")
    cons=""
    for element in cons_element:
        cons=cons+element.text
    cons=cons.replace("Cons\n","")
    report.add_heading("Cons",level=2)
    report.add_paragraph(cons)
    #pros
    pros_element = soup.findAll('div', class_="pros")
    pros=""
    for element in pros_element:
        pros=pros+element.text
    pros=pros.replace("Pros\n","")
    report.add_heading("Pros",level=2)
    report.add_paragraph(pros)
    #Sales,OperatingProfit & NetProfits and Cash Flow
    tables=pandas.read_html("/home/yashas/Desktop/Quant/NEW_DATA/"+code+".html")
    PandL=tables[1].copy()
    PandL=PandL.transpose()
    x=PandL.index
    x=x[-5:]
    sales=PandL[0].copy()
    operating_profit=PandL[2].copy()
    net_profit=PandL[9].copy()
    sales=list(sales)
    sales=sales[1:]
    sales_num=[]
    for i in sales:
        sales_num.append(tools.convert_to_number(i))
    operating_profit=list(operating_profit)
    operating_profit=operating_profit[1:]
    operating_profit_num=[]
    for i in operating_profit:
        operating_profit_num.append(tools.convert_to_number(i))
    net_profit=list(net_profit)
    net_profit=net_profit[1:]
    net_profit_num=[]
    for i in net_profit:
        net_profit_num.append(tools.convert_to_number(i))
    cash_flow=tables[7].copy()
    cash_flow=cash_flow.transpose()
    net_cash_flow=cash_flow[3].copy()
    net_cash_flow=list(net_cash_flow)
    net_cash_flow_num=[]
    net_cash_flow=net_cash_flow[1:]
    for i in net_cash_flow:
        net_cash_flow_num.append(tools.convert_to_number(i))
    plt.plot(x,net_cash_flow_num[-5:],label="Net Cash Flow")
    plt.plot(x,sales_num[-5:],label="Sales")
    plt.plot(x,operating_profit_num[-5:],label="Operating Profit")
    plt.plot(x,net_profit_num[-5:],label="Net Profit")
    plt.legend()
    plt.xlabel("Year")
    plt.ylabel("Amount")
    plt.title("Some Numbers(in Rs Crores)")
    plt.savefig("/home/yashas/Desktop/Quant/report/"+code+".png")
    report.add_heading("Sales,Operating Profit, Net Profits and Net Cash Flow",level=2)
    report.add_picture("/home/yashas/Desktop/Quant/report/"+code+".png",width=Inches(4.5))
    plt.close()
    '''
    #Cash Flows
    cash_flow=tables[7].copy()
    cash_flow=cash_flow.transpose()
    x=cash_flow.index
    x=x[-5:]
    net_cash_flow=cash_flow[3].copy()
    net_cash_flow=list(net_cash_flow)
    net_cash_flow_num=[]
    net_cash_flow=net_cash_flow[1:]
    for i in net_cash_flow:
        net_cash_flow_num.append(tools.convert_to_number(i))
    plt.plot(x,net_cash_flow_num[-5:],label="Net Cash Flow")
    plt.xlabel("Year")
    plt.ylabel("Amount")
    plt.legend()
    plt.title("Net Cash Flow")
    plt.savefig("/home/yashas/Desktop/Quant/report/"+code+".png")
    report.add_heading("Cash Flows",level=2)
    report.add_picture("/home/yashas/Desktop/Quant/report/"+code+".png",width=Inches(5.5))
    plt.close()'''
    #Shareholding Pattern
    try:
        shareholding_pattern=tables[9].copy()
        shareholding_pattern=shareholding_pattern.transpose()
        lables=shareholding_pattern.iloc[0]
        lables=list(lables)
        temp=[]
        for i in lables:
            temp.append(i.replace("\xa0+",""))
        lables=temp.copy()
        shareholding_pattern=shareholding_pattern.iloc[-1]
        shareholding_pattern=list(shareholding_pattern)
        shareholding_pattern_num=[]
        for i in shareholding_pattern:
            shareholding_pattern_num.append(tools.convert_to_number(i))
        plt.figure(figsize=(6,6))
        plt.pie(shareholding_pattern_num,labels=lables,autopct='%1.1f%%',startangle=140)
        plt.title("Shareholding Pattern")
        plt.savefig("/home/yashas/Desktop/Quant/report/"+code+".png")
        report.add_heading("Shareholding Pattern",level=2)  
        report.add_picture("/home/yashas/Desktop/Quant/report/"+code+".png",width=Inches(4.5))
        plt.close()
    except Exception as exe:
        print(exe)
        traceback.print_exc()
        print("No Shareholding data found")
        report.add_paragraph("No Shareholding data found")
    #report.save("/home/yashas/Desktop/Quant/report/"+code+".docx")
    #print(f"Report for {company_name} saved")